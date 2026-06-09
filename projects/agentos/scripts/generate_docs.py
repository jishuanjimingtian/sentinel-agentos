"""
AgentOS 技术文档生成器
生成三份 Word 文档：设计文档、架构文档、任务清单
"""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from datetime import date

OUT_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "docs")

def set_cell_shading(cell, color):
    """设置单元格底色"""
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def style_doc(doc):
    """统一文档样式"""
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.25

    for level in range(1, 4):
        h_style = doc.styles[f'Heading {level}']
        h_style.font.name = '微软雅黑'
        if level == 1:
            h_style.font.size = Pt(22)
            h_style.font.bold = True
            h_style.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
        elif level == 2:
            h_style.font.size = Pt(16)
            h_style.font.bold = True
            h_style.font.color.rgb = RGBColor(0x2d, 0x2d, 0x3f)
        elif level == 3:
            h_style.font.size = Pt(13)
            h_style.font.bold = True
            h_style.font.color.rgb = RGBColor(0x3d, 0x3d, 0x5c)

    # 代码块样式
    if 'Code' not in [s.name for s in doc.styles]:
        code_style = doc.styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
        code_style.font.name = 'Consolas'
        code_style.font.size = Pt(9)
        code_style.paragraph_format.space_before = Pt(4)
        code_style.paragraph_format.space_after = Pt(4)
        code_style.paragraph_format.left_indent = Cm(1)


def add_table(doc, headers, rows, col_widths=None):
    """添加格式化表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)

    # Rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph()
    return table


def add_code_block(doc, code_text):
    """添加代码块"""
    for line in code_text.strip().split('\n'):
        p = doc.add_paragraph(line, style='Code')
        for run in p.runs:
            run.font.name = 'Consolas'
            run.font.size = Pt(9)


# ============================================================
# 文档 1: 设计文档 (DESIGN.docx)
# ============================================================
def generate_design_doc():
    doc = Document()
    style_doc(doc)

    # 封面
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('AgentOS 设计文档')
    run.font.size = Pt(32)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Agent Runtime Infrastructure — 让任何 Agent 变得可靠、可审计、可改进')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x88)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f'\n版本：v0.1.0-draft\n').font.size = Pt(11)
    meta.add_run(f'状态：设计阶段\n').font.size = Pt(11)
    meta.add_run(f'日期：{date.today().strftime("%Y-%m-%d")}\n').font.size = Pt(11)
    meta.add_run('作者：安妮 (Annie) @ AgentOS Team').font.size = Pt(11)

    doc.add_page_break()

    # ====== 1. 产品定位 ======
    doc.add_heading('1. 产品定位', level=1)

    doc.add_heading('1.1 一句话描述', level=2)
    doc.add_paragraph(
        'AgentOS = 确定性 Guard 层 + 分层记忆 + 自动评估，'
        '让任何 Agent 变得可靠、可审计、可改进。'
    )

    doc.add_heading('1.2 不是 Agent，是 Agent 的操作系统', level=2)
    add_table(doc, ['操作系统层', '对应 AgentOS 模块', '说明'], [
        ['应用程序', '任意 Agent', 'OpenClaw / LangChain / CrewAI / 自研'],
        ['操作系统内核', 'Schema Gate + Risk Gate', '确定性代码，零 LLM 依赖'],
        ['文件系统', '分层 Memory Store', 'Working / Episodic / Semantic'],
        ['日志系统', 'Audit Log', '操作级审计（不可篡改，支持回滚）'],
        ['性能监控', 'Evaluator', '三阶段评估 + 隐性反馈捕获'],
    ])

    doc.add_heading('1.3 解决什么', level=2)
    add_table(doc, ['痛点', 'AgentOS 方案', '核心思路'], [
        ['Agent 幻觉导致错误操作', 'Schema Gate + Verify Gate', '确定性校验，不依赖 LLM 判断'],
        ['Agent 越权/危险操作', 'Risk Gate（四维风险评分）', '纯数学公式，0-100 自动分级'],
        ['Agent 记不住、记不对', 'Memory 层事件驱动自动记录', '分层记忆架构，不是 RAG'],
        ['出事了查不到原因', '操作级审计日志', '每次 tool call 前后 diff'],
        ['不知道 Agent 好不好', '三阶段评估 + 隐性反馈', '不止看结果，看全过程'],
    ])

    doc.add_heading('1.4 目标用户', level=2)
    add_table(doc, ['阶段', '用户', '场景'], [
        ['MVP', '独立开发者', '自部署 SDK，接入自己的 Agent'],
        ['v1.0', '小型团队', '团队级 Dashboard，共享 Guard Rules'],
        ['v2.0', '企业', 'SaaS 云端，企业级审计合规'],
    ])

    # ====== 2. 核心问题分析 ======
    doc.add_heading('2. 核心问题分析', level=1)

    doc.add_heading('2.1 为什么现有方案都不行', level=2)

    doc.add_heading('LangChain / LlamaIndex 的 Memory', level=3)
    doc.add_paragraph('本质是「把对话历史扔进向量库」— 只有检索，没有理解', style='List Bullet')
    doc.add_paragraph('缺乏分层：所有记忆一视同仁，不区分重要性', style='List Bullet')
    doc.add_paragraph('无遗忘机制：记忆无限膨胀，检索精度指数级下降', style='List Bullet')
    doc.add_paragraph('Agent 需手动管理记忆 → 又回到「agent 自己能记住」的悖论', style='List Bullet')

    doc.add_heading('CrewAI / AutoGPT 的安全', level=3)
    doc.add_paragraph('安全靠 prompt："请不要删除重要文件" — 这是愿望，不是约束', style='List Bullet')
    doc.add_paragraph('无审计：agent 做了什么、为什么做、做对了没，全不可追溯', style='List Bullet')
    doc.add_paragraph('无回滚：出事了只能手动 git reset', style='List Bullet')

    doc.add_heading('所有 Agent 框架的评估', level=3)
    doc.add_paragraph('不存在。没有框架做 agent 的自动化评估', style='List Bullet')
    doc.add_paragraph('最多有个 "success rate" 计数器，不分析为什么会失败', style='List Bullet')
    doc.add_paragraph('完全不捕捉用户的隐性反馈信号', style='List Bullet')

    doc.add_heading('2.2 为什么 Guard 层是地基', level=2)
    add_table(doc, ['痛点', 'Guard 如何解决', '关键'], [
        ['可靠性（痛点一）', '拦截幻觉操作', '前端过滤，不靠 LLM'],
        ['工具调用（痛点二）', 'Schema Gate 校验参数', '格式错误第一步拦截'],
        ['记忆（痛点三）', 'Guard 日志 = 精确事件源', '记忆建立在准确数据上'],
        ['安全（痛点四）', 'Risk Gate 分级处置', '从愿望变成物理约束'],
        ['评估（痛点五）', 'Guard 日志 + Verify 结果', '评估有真实数据基础'],
    ])

    # ====== 3. 设计哲学 ======
    doc.add_heading('3. 设计哲学', level=1)

    doc.add_heading('3.1 三条铁律', level=2)
    add_table(doc, ['原则', '含义', '反例'], [
        ['确定性优先', '能不用 LLM 就不用的功能，必须用确定性代码实现', '用 LLM 做安全判断'],
        ['可审计性优先', '所有操作必须可追溯、可回滚、可解释', 'agent 操作后无日志，出事查不到'],
        ['渐进增强', '框架无关，可增量接入。不要求替换现有 Agent 架构', 'LangChain 的强绑定'],
    ])

    doc.add_heading('3.2 边界定义', level=2)
    doc.add_heading('AgentOS 不做', level=3)
    doc.add_paragraph('不替代 Agent 框架（OpenClaw / LangChain / CrewAI）', style='List Bullet')
    doc.add_paragraph('不替代 LLM（不训练模型）', style='List Bullet')
    doc.add_paragraph('不替代编排（不决定 agent 做什么任务）', style='List Bullet')
    doc.add_paragraph('不提供 UI（SDK + API，不做前端）', style='List Bullet')

    doc.add_heading('AgentOS 只做', level=3)
    doc.add_paragraph('拦截 tool call → 校验 + 分级', style='List Bullet')
    doc.add_paragraph('记录操作 → 事件采集 + 分层存储', style='List Bullet')
    doc.add_paragraph('评估质量 → 收集指标 + 隐性反馈', style='List Bullet')
    doc.add_paragraph('暴露 API → 查询、回滚、分析', style='List Bullet')

    # ====== 4. Guard 层设计 ======
    doc.add_heading('4. Guard 层设计', level=1)

    doc.add_heading('4.1 整体流水线', level=2)
    doc.add_paragraph(
        'Guard 层由 6 个阶段组成一条线性流水线。每个阶段只做一件事，'
        '前两个阶段（Schema Gate + Risk Gate）是纯确定性的，零 LLM 依赖。'
    )
    add_table(doc, ['阶段', '名称', '职责', '依赖 LLM?', '失败处理'], [
        ['1', 'Schema Gate', '校验参数格式和约束', '否（纯 JSON Schema）', '拒绝执行'],
        ['2', 'Risk Gate', '计算风险评分，分级处置', '否（纯数学公式）', '通知/确认/拒绝'],
        ['3', 'Snapshot', '执行前快照', '否', '记录警告但继续'],
        ['4', 'Execute', '执行 tool call', '否', '记录错误'],
        ['5', 'Verify Gate', '执行后校验结果', '否', '触发回滚'],
        ['6', 'Audit Log', '写不可篡改日志', '否', '系统错误'],
    ])

    doc.add_heading('4.2 Schema Gate（第一关）', level=2)
    doc.add_paragraph(
        'Schema Gate 是 Guard 层的第一道防线。校验 tool call 参数的格式和约束。'
        '100% 确定性——本质是 JSON Schema 验证 + 自定义路径约束，不需要任何 AI 判断。'
    )

    doc.add_heading('校验项', level=3)
    add_table(doc, ['校验项', '实现方式', '示例'], [
        ['必填参数', 'JSON Schema required', 'delete_file 必须提供 path'],
        ['参数类型', 'JSON Schema type', 'path 必须是 string'],
        ['数值范围', 'JSON Schema min/max', 'max_tokens 在 1-100000'],
        ['枚举值', 'JSON Schema enum', 'mode 只能是 read/write/append'],
        ['路径约束', '自定义 validator', 'path 必须在 workspace 内'],
        ['正则约束', 'JSON Schema pattern', 'branch_name 必须符合命名规则'],
        ['依赖约束', '自定义 validator', 'auto_merge=true → base_branch 必填'],
        ['互斥约束', '自定义 validator', 'content 和 file_path 不能同时存在'],
    ])

    doc.add_heading('扩展 Schema 格式', level=3)
    doc.add_paragraph(
        '在标准 JSON Schema 基础上，引入 x- 前缀的自定义扩展字段来表达工具语义约束。'
    )
    add_table(doc, ['扩展字段', '用途', '示例值'], [
        ['x-path-scope', '路径约束范围', 'workspace / temp / global'],
        ['x-path-allow', '允许的路径 Glob 模式', '["*.ts", "src/**"]'],
        ['x-path-deny', '禁止的路径 Glob 模式', '[".git/**", "*.key"]'],
        ['x-max-size', '参数值最大字节数', '1048576'],
        ['x-secret', '标记为敏感参数（日志脱敏）', 'true'],
        ['x-depends-on', '参数间依赖关系', '{"auto_merge": ["base_branch"]}'],
        ['x-mutually-exclusive', '互斥参数组', '[["content", "file_path"]]'],
    ])

    doc.add_heading('4.3 Risk Gate（第二关）', level=2)
    doc.add_paragraph(
        'Risk Gate 用纯数学公式计算操作的风险评分。'
        '不依赖 LLM 判断——因为风险评估可以建模为四个独立维度的乘积。'
    )

    doc.add_heading('风险公式', level=3)
    p = doc.add_paragraph()
    run = p.add_run('RiskScore = Impact × Irreversibility × Sensitivity × (1 + ErrorRate)')
    run.bold = True
    run.font.size = Pt(12)

    doc.add_heading('Impact（影响范围）', level=3)
    add_table(doc, ['级别', '分值', '定义', '示例'], [
        ['local', 1, '只影响单个文件或内存', '读文件、搜索、内存计算'],
        ['workspace', 3, '影响工作区文件', '写文件、git commit'],
        ['project', 6, '影响整个项目', 'npm publish、git push、删除目录'],
        ['system', 10, '影响系统或外部', '修改系统配置、sudo、发邮件'],
    ])

    doc.add_heading('Irreversibility（可逆程度）', level=3)
    add_table(doc, ['分值', '含义', '示例'], [
        [1.0, '完全可逆', '读文件、搜索'],
        [0.8, '可逆但有轻微残留', '写文件（git 可回滚）'],
        [0.5, '半可逆，需手动操作', 'git push（force push 影响他人）'],
        [0.2, '基本不可逆', 'npm publish（unpublish 有时间窗口）'],
        [0.0, '完全不可逆', '删除远程数据库、发送 email'],
    ])

    doc.add_heading('Sensitivity（数据敏感度）', level=3)
    add_table(doc, ['分值', '含义', '示例'], [
        [0.0, '非敏感', '公开代码、README、日志'],
        [0.3, '低敏感', '项目配置、构建脚本'],
        [0.6, '中敏感', '业务逻辑代码、API 配置'],
        [0.9, '高敏感', '.env、密钥、用户数据'],
        [1.0, '极高敏感', '生产环境凭证、客户 PII'],
    ])

    doc.add_heading('ErrorRate（历史错误率）', level=3)
    doc.add_paragraph('ErrorRate = 该工具过去失败次数 / 该工具过去总调用次数')
    doc.add_paragraph('初始值（冷启动）：只读类 0.01、写入类 0.05、删除类 0.10、网络类 0.08')
    doc.add_paragraph('随使用动态更新。错误率上升 → 风险分数上升 → 更多操作需要确认。')

    doc.add_heading('风险分数示例', level=3)
    add_table(doc, ['操作', 'Impact', 'Rev.', 'Sens.', 'Err', '分数', '动作'], [
        ['read_file', '1', '1.0', '0.0', '0.01', '0.00', '放行'],
        ['write_file(src/*.ts)', '3', '0.8', '0.3', '0.05', '0.19', '放行'],
        ['git_push', '6', '0.5', '0.3', '0.04', '0.94', '通知'],
        ['npm_publish', '8', '0.2', '0.5', '0.06', '3.39', '确认'],
        ['delete_dir(workspace)', '6', '0.2', '0.6', '0.08', '3.11', '确认'],
        ['shell_rm_rf', '10', '0.0', '0.9', '0.02', '9.18', '拒绝'],
        ['delete_prod_db', '10', '0.0', '1.0', '0.01', '10.10', '拒绝'],
    ])

    doc.add_heading('可配置阈值', level=3)
    add_table(doc, ['阈值参数', '默认值', '含义'], [
        ['auto_approve', '0.5', '≤ 0.5 自动放行'],
        ['notify', '1.0', '≤ 1.0 执行后通知用户'],
        ['confirm', '3.0', '≤ 3.0 暂停等待用户确认'],
        ['deny', '8.0', '> 8.0 直接拒绝'],
    ])

    # ====== 4.4 Snapshot ======
    doc.add_heading('4.4 Snapshot（执行前快照）', level=2)
    doc.add_paragraph(
        'Snapshot 在执行前记录关键状态的 hash，为 Verify 和 Rollback 提供基准。'
        '不做全量备份（太重），只做变更追踪级别的快照。'
    )
    add_table(doc, ['场景', '策略', '原因'], [
        ['单文件修改', '只 hash 被修改的文件', '节省时间'],
        ['批量修改', 'hash 所有涉及文件', '需要知道改了哪些'],
        ['危险操作 (>3.0)', 'hash 整个 workspace', '确保完整回滚'],
        ['删除操作', 'hash 目标 + 记录完整内容', '回滚 = 恢复内容'],
    ])

    doc.add_heading('4.5 Verify Gate（第五关）', level=2)
    doc.add_paragraph(
        '执行后校验结果，检测幻觉和错误。所有校验项都是确定性的——不需要 LLM 判断。'
    )
    add_table(doc, ['校验项', '触发条件', '判定标准', '失败处理'], [
        ['文件存在性', '声称创建/修改了文件', 'fs.existsSync', '标记 FAIL（幻觉！）'],
        ['文件变更', '声称修改了文件', 'hash 对比 Snapshot', '未变更→FAIL'],
        ['Lint 通过', '修改了代码文件', 'eslint --quiet', '未通过→WARN'],
        ['TypeCheck', '修改了 .ts 文件', 'tsc --noEmit', '未通过→WARN'],
        ['格式合法', '声称返回 JSON', 'JSON.parse', '格式错误→FAIL'],
        ['返回值非空', '应返回内容', 'result.length > 0', '空结果→WARN'],
        ['npm 发布验证', '声称发布了 npm 包', 'npm view', '未发布→FAIL（幻觉！）'],
        ['git push 验证', '声称推送成功', 'git ls-remote', '未推送→FAIL'],
    ])

    doc.add_heading('4.6 Audit Log（第六关）', level=2)
    doc.add_paragraph(
        '记录每一次操作的完整信息：什么时间、哪个 agent、做了什么、'
        '结果是什么、执行前后 diff。追加写入，不可篡改。'
    )
    add_table(doc, ['版本', '方案', '适用'], [
        ['MVP', '追加写入 JSONL 文件，文件设置为只读', '单机部署'],
        ['v1.0', '写入本地 SQLite，WAL 模式，无 DELETE 权限', '团队使用'],
        ['v2.0', '远端 append-only log service', '企业级审计'],
    ])

    doc.add_heading('4.7 Rollback（回滚）', level=2)
    doc.add_paragraph('触发条件：Verify Gate 返回 FAIL 且 Risk > confirm 阈值，或用户手动触发。')
    doc.add_paragraph('回滚过程：')
    doc.add_paragraph('1. 从 Audit Log 找到该操作的 Snapshot', style='List Number')
    doc.add_paragraph('2. 对比 Snapshot 与当前文件状态', style='List Number')
    doc.add_paragraph('3. 对每个被修改的文件：git checkout 或从 Snapshot 恢复', style='List Number')
    doc.add_paragraph('4. 记录回滚日志，返回回滚报告', style='List Number')

    # ====== 5. Memory 层 ======
    doc.add_heading('5. Memory 层设计', level=1)

    doc.add_heading('5.1 核心理念：像人脑一样记忆', level=2)
    add_table(doc, ['人类记忆', '功能', 'AgentOS 对应层', '存活周期', '大小上限'], [
        ['工作记忆', '当前对话、几秒内的事', 'Working Memory', '1 session', '< 50KB'],
        ['情景记忆', '昨天发生了什么、关键决策', 'Episodic Memory', '数周-数月', '< 500KB'],
        ['语义记忆', '学到的知识、规则、偏好', 'Semantic Memory', '永久', '< 100KB'],
    ])

    doc.add_heading('5.2 Working Memory（工作记忆）', level=2)
    doc.add_paragraph('用途：当前 session 的实时上下文。session 结束时自动清空。')
    doc.add_paragraph('数据来源：Guard 层的 Audit Log（自动）、LLM 输出（自动）、用户输入（自动）。')
    doc.add_paragraph('关键行为：')
    doc.add_paragraph('session 结束时自动清空，重要信息在清空前自动升级到 Episodic Memory', style='List Bullet')
    doc.add_paragraph('超过上下文预算时自动压缩旧消息', style='List Bullet')

    doc.add_heading('5.3 Episodic Memory（情景记忆）', level=2)
    doc.add_paragraph('用途：跨 session 的事件记录——"上周三发生了什么"。')
    doc.add_paragraph('事件类型：')
    add_table(doc, ['事件类型', '重要性基数', '说明'], [
        ['correction', 0.9, '人类纠正（最高优先级，永不过期）'],
        ['decision', 0.8, '人类做出决策'],
        ['milestone', 0.8, '里程碑事件'],
        ['user_feedback', 0.8, '用户显式反馈'],
        ['publish', 0.7, '发布/部署'],
        ['error', 0.7, '系统错误'],
        ['tool_failure', 0.6, 'tool call 失败'],
        ['note', 0.3, 'agent 主动记录备注'],
        ['tool_call', 0.2, '常规 tool call'],
    ])

    doc.add_heading('重要性自动评分公式', level=3)
    p = doc.add_paragraph()
    run = p.add_run('Importance = BaseImportance × RecencyBoost × FrequencyBoost × FeedbackBoost')
    run.bold = True
    doc.add_paragraph('RecencyBoost: 越新越重要，30天后衰减至 1.0x（不再加成）')
    doc.add_paragraph('FrequencyBoost: 同一标签每次重复 +0.1，最多 +0.5')
    doc.add_paragraph('FeedbackBoost: 用户说"记住这个"→ +0.3；用户纠正 → +0.5')

    doc.add_heading('压缩策略', level=3)
    add_table(doc, ['阶段', '触发条件', '内容'], [
        ['full → summary', '重要性 < 0.3 且 > 7天', '调用 LLM 摘要'],
        ['summary → one-liner', '重要性 < 0.2 且 > 30天', '压缩为一句话'],
        ['one-liner → forgotten', '重要性 < 0.1 且 > 90天', '删除'],
        ['永不压缩', 'importance > 0.7', '永久保留 full 详情'],
    ])

    doc.add_heading('5.4 Semantic Memory（语义记忆）', level=2)
    doc.add_paragraph('用途：从情景记忆中提炼出的持久知识，永不衰减。')
    add_table(doc, ['分类', '内容', '更新规则'], [
        ['用户偏好', 'key-value 偏好字典', '同一偏好表达 ≥2 次 → 写入'],
        ['项目上下文', '技术栈、编码规范、架构、已知问题', 'agent 主动识别 + 用户确认'],
        ['学到的规则', '"发布 npm 前必须更新 CHANGELOG"', '同类 correction ≥3 次 → 提炼规则'],
        ['术语表', '项目特定术语 → 含义', '新术语出现 → 主动询问确认'],
    ])

    doc.add_heading('5.5 Session 启动上下文注入', level=2)
    doc.add_paragraph(
        'AgentOS 的核心价值之一：新 session 启动时自动注入最相关的上下文。'
    )
    doc.add_paragraph('从 Semantic Memory 加载：用户偏好、项目信息、最近学到的规则', style='List Bullet')
    doc.add_paragraph('从 Episodic Memory 加载：最近 7 天高重要性事件、最近 milestone、最近 correction', style='List Bullet')
    doc.add_paragraph('自动生成上下文摘要注入到 system prompt，agent 无需手动查询记忆', style='List Bullet')

    # ====== 6. Evaluator ======
    doc.add_heading('6. Evaluator 层设计', level=1)

    doc.add_heading('6.1 核心理念', level=2)
    doc.add_paragraph(
        '大多数 Agent 框架的"评估"是「做完任务，人工看结果」。'
        'AgentOS 的评估是「每个操作自动打分，累积成 agent 的质量画像」。'
    )

    doc.add_heading('6.2 三阶段评估', level=2)
    add_table(doc, ['阶段', '评估时机', '关键指标'], [
        ['Pre-exec', '执行前', 'Schema 通过率、风险分布、参数质量'],
        ['Runtime', '执行中', '重试率、自我修正率、超时率、工具选择准确度'],
        ['Post-exec', '执行后', 'Verify 通过率、用户接受度、结果利用度、任务完成度'],
    ])

    doc.add_heading('6.3 隐性反馈捕获（核心差异点）', level=2)
    doc.add_paragraph('为什么隐性反馈重要：')
    doc.add_paragraph('用户说"好的谢谢"可能是礼貌 → 不代表真的满意', style='List Bullet')
    doc.add_paragraph('用户修改了 agent 的输出 → 精确的不满意信号（但很少被捕获）', style='List Bullet')
    doc.add_paragraph('用户重复相同指令 → agent 第一次没做对（强烈的负反馈）', style='List Bullet')
    doc.add_paragraph('用户秒回"不对" → 错误明显（但为时已晚）', style='List Bullet')

    add_table(doc, ['信号类型', '信号强度', '含义', '如何捕获'], [
        ['accept_as_is', '+0.8', '用户直接使用输出', '检测用户下一个操作不是修改该输出'],
        ['accept_with_edit', '-0.3', '用户修改后使用', '检测用户编辑了 agent 创建的文件'],
        ['reject_and_redo', '-0.7', '用户删掉重做', '检测 agent 输出被删除+用户重新发指令'],
        ['repeat_instruction', '-0.9', '用户重复相同指令', '检测语义相同的新指令'],
        ['silent_proceed', '+0.3', '用户继续新话题', '用户立即发新指令（说明上一个满意）'],
        ['explicit_correct', '-1.0', '用户说"不对"', '文本匹配'],
    ])

    doc.add_heading('6.4 评估 Dashboard（v1.0）', level=2)
    add_table(doc, ['模块', '展示内容'], [
        ['可靠性趋势', 'Verify 通过率随时间变化（应该上升）'],
        ['安全趋势', 'Risk 分数分布、被拒绝的操作占比'],
        ['效率趋势', '平均重试次数、自我修正率、操作耗时'],
        ['用户满意度', '隐性反馈综合评分（接受率 vs 修改率）'],
        ['Cost 趋势', '每条高质量操作的 token 成本（应该下降）'],
    ])

    # ====== 7. API 设计 ======
    doc.add_heading('7. API 设计', level=1)

    doc.add_heading('7.1 接入方式：5 行代码', level=2)
    add_code_block(doc, """import { AgentOS } from '@agentos/core';

const aos = new AgentOS({
  memory: { backend: 'local' },
  guard: { rules: './guard-rules.yaml' },
  eval: { track: true }
});

// 包裹任意 Agent 的 tool executor
const safeExecutor = aos.wrap(originalExecutor);
// 现在所有 tool call 都有 记忆 + 安全 + 评估""")

    doc.add_heading('7.2 核心 API', level=2)
    add_table(doc, ['模块', 'API', '说明'], [
        ['Guard', 'aos.guard.intercept(toolCall)', '拦截 tool call → 返回 GuardResult'],
        ['Guard', 'aos.guard.snapshot(scope)', '手动创建快照'],
        ['Guard', 'aos.guard.rollback(snapshotId)', '回滚到指定快照'],
        ['Memory', 'aos.memory.recall(query)', '语义 + 时间双维度检索'],
        ['Memory', 'aos.memory.context()', '获取当前 session 的上下文摘要'],
        ['Memory', 'aos.memory.remember(event)', '手动记录事件'],
        ['Eval', 'aos.eval.report(sessionId?)', '获取评估报告'],
        ['Eval', 'aos.eval.feedback(signal)', '记录显式/隐式反馈'],
        ['Core', 'aos.wrap(executor)', '包裹 tool executor，一行接入'],
    ])

    # ====== 8. 配置 ======
    doc.add_heading('8. 配置系统', level=1)

    doc.add_heading('8.1 默认配置', level=2)
    add_code_block(doc, """# agentos.yaml
guard:
  risk:
    auto_approve: 0.5
    notify: 1.0
    confirm: 3.0
    deny: 8.0
    
  verify:
    lint: true
    typecheck: true
    npm_verify: true
    
  audit:
    backend: jsonl
    max_size_mb: 100

memory:
  working:
    max_messages: 50
    max_tokens: 10000
    
  episodic:
    max_size_kb: 500
    compress_after_days: 7
    forget_after_days: 90
    
  session_inject:
    recent_days: 7
    max_events: 20
    max_rules: 10

eval:
  track_implicit: true
  report_interval_hours: 24""")

    doc.add_heading('8.2 工具注册（扩展 Schema）', level=2)
    add_code_block(doc, """// tools/delete_file.tool.json
{
  "name": "delete_file",
  "description": "Delete a file at the specified path",
  "parameters": {
    "type": "object",
    "properties": {
      "path": {
        "type": "string",
        "x-path-scope": "workspace",
        "x-path-deny": [".git/**", "*.key", "*.pem", ".env"]
      }
    },
    "required": ["path"]
  },
  "risk": {
    "impact": "workspace",
    "reversibility": 0.3,
    "sensitivity": 0.3,
    "initialErrorRate": 0.05
  },
  "verify": {
    "postCheck": "file_deleted"
  }
}""")

    # ====== 9. 数据模型 ======
    doc.add_heading('9. 数据模型', level=1)

    doc.add_heading('9.1 核心实体关系', level=2)
    add_table(doc, ['实体', '关键字段', '生命周期'], [
        ['Session', 'id, agentId, startedAt, status', '从创建到关闭'],
        ['ToolCall', 'id, sessionId, toolName, params, status', '每次 tool call'],
        ['AuditEntry', 'id, toolCallId, schemaGate, riskGate, verifyGate, diff, rollback', '永久（追加）'],
        ['Snapshot', 'id, toolCallId, fileHashes, gitHead', '关联 AuditEntry 生命周期'],
        ['EpisodicEvent', 'id, type, importance, compression, content, tags', '根据重要性压缩或遗忘'],
        ['SemanticMemory', 'userPreferences, projectContext, learnedRules', '永久，持续更新'],
        ['GuardRule', 'toolName, riskParams, verifyChecks', '随工具注册'],
    ])

    doc.add_heading('9.2 存储后端', level=2)
    add_table(doc, ['后端', '适用场景', '存储内容'], [
        ['JSONL 文件', 'MVP 单机', 'Audit Log'],
        ['SQLite', 'v1.0 团队', 'Audit + Episodic + Semantic'],
        ['远端 Log Service', 'v2.0 企业', 'Audit（append-only）'],
        ['内存', 'Runtime', 'Working Memory'],
    ])

    # ====== 保存 ======
    doc.save(os.path.join(OUT_DIR, 'AgentOS_设计文档.docx'))
    print('[OK] AgentOS_设计文档.docx')


# ============================================================
# 文档 2: 架构文档 (ARCHITECTURE.docx)
# ============================================================
def generate_architecture_doc():
    doc = Document()
    style_doc(doc)

    # 封面
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('AgentOS 架构文档')
    run.font.size = Pt(32)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('模块架构 · 数据流 · 部署方案')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x88)

    doc.add_page_break()

    # ====== 1. 总览 ======
    doc.add_heading('1. 系统总体架构', level=1)
    doc.add_paragraph(
        'AgentOS 采用分层架构：上层是 Agent 框架（任意），下层是 AgentOS Runtime。'
        'AgentOS 本身分三大模块（Guard / Memory / Evaluator），通过 SDK 或 MCP Server 暴露给 Agent。'
    )

    add_table(doc, ['层', '职责', '关键技术'], [
        ['Agent 框架层', '任意 Agent（OpenClaw / LangChain / CrewAI）', '业务逻辑'],
        ['SDK / MCP Server', 'AgentOS 对外接口', 'TypeScript + MCP 协议'],
        ['Guard 层', 'Tool call 拦截、校验、分级、审计', 'JSON Schema + 确定性规则引擎'],
        ['Memory 层', '三层记忆存储与查询', 'SQLite / JSONL + 向量检索'],
        ['Evaluator 层', '三阶段评估 + 隐性反馈', '统计分析 + 信号检测'],
        ['存储层', '数据持久化', 'JSONL / SQLite / 远端 Log Service'],
    ])

    # ====== 2. 模块架构 ======
    doc.add_heading('2. 模块架构详解', level=1)

    doc.add_heading('2.1 Guard 模块', level=2)
    doc.add_paragraph('Guard 模块是一条 6 阶段线性流水线。每个阶段有明确的输入、处理、输出。')

    add_table(doc, ['阶段', '输入', '处理', '输出', '是否阻塞'], [
        ['Schema Gate', 'ToolCall 请求', 'JSON Schema + x- 扩展校验', 'GateResult (pass/fail)', '是（失败→拒绝）'],
        ['Risk Gate', '通过 Schema 的 ToolCall', '四维风险公式计算', 'RiskScore + Action', '否（但影响后续流程）'],
        ['Snapshot', '将被执行的操作 + 环境', 'hash 相关文件/状态', 'Snapshot 对象', '否'],
        ['Execute', 'ToolCall + Risk Action', '执行（direct/sandbox/dry-run）', 'ToolCall 结果', '否'],
        ['Verify Gate', 'ToolCall 结果 + Snapshot', '确定性校验项检查', 'VerifyResult (PASS/WARN/FAIL)', '否（FAIL→回滚）'],
        ['Audit Log', '所有前述阶段的结果', '序列化 + 追加写入', '持久化日志条目', '否'],
    ])

    doc.add_heading('2.2 Memory 模块', level=2)
    doc.add_paragraph('Memory 模块由三个子模块和一个管理器组成。')

    add_table(doc, ['子模块', '触发时机', '数据流向'], [
        ['Working Memory (WM)', '每次对话/操作 → 自动写入', 'Guard AuditLog → WM'],
        ['Episodic Memory (EM)', 'Session 结束 / 定期压缩', 'WM 重要事件 → EM；EM 旧事件 → 压缩'],
        ['Semantic Memory (SM)', '定期提炼 / 用户显式', 'EM 高频 pattern → SM 规则；用户输入 → SM 偏好'],
        ['Memory Manager', 'Session 启动 / 查询时', 'SM + EM → 上下文摘要 → 注入 System Prompt'],
    ])

    doc.add_heading('2.3 Evaluator 模块', level=2)
    doc.add_paragraph('Evaluator 持续收集指标，生成 Agent 质量画像。')

    add_table(doc, ['子模块', '数据源', '产出'], [
        ['Pre-exec Collector', 'Schema Gate 结果', 'schema_pass_rate, risk_distribution'],
        ['Runtime Collector', 'Execute 过程', 'retry_rate, self_correction, timeout_rate'],
        ['Post-exec Collector', 'Verify Gate + 用户行为', 'verify_pass_rate, implicit_feedback'],
        ['Report Generator', '所有 Collector 数据', 'Agent Quality Score + 趋势图数据'],
    ])

    # ====== 3. 数据流 ======
    doc.add_heading('3. 核心数据流', level=1)

    doc.add_heading('3.1 Tool Call 完整数据流', level=2)
    doc.add_paragraph('下图描述一次 tool call 从请求到完成的完整数据流：')
    doc.add_paragraph('1. Agent → SDK/MCP Server: 发出 ToolCall 请求', style='List Number')
    doc.add_paragraph('2. SDK → Guard: intercept(toolCall)', style='List Number')
    doc.add_paragraph('3. Guard → Schema Gate: 校验参数格式。Fail → 返回错误给 Agent', style='List Number')
    doc.add_paragraph('4. Guard → Risk Gate: 计算风险评分 → 根据阈值决定 auto/notify/confirm/deny', style='List Number')
    doc.add_paragraph('5. 如果 confirm → 暂停流水线，发送确认请求给用户。超时或拒绝 → 返回给 Agent', style='List Number')
    doc.add_paragraph('6. Guard → Snapshot: 快照当前状态', style='List Number')
    doc.add_paragraph('7. Guard → Execute: 执行 tool call（可能 sandbox/dry-run）', style='List Number')
    doc.add_paragraph('8. Guard → Verify Gate: 校验执行结果。FAIL → 触发回滚', style='List Number')
    doc.add_paragraph('9. Guard → Audit Log: 写入不可篡改日志', style='List Number')
    doc.add_paragraph('10. Guard → Memory: Audit Log 自动写入 Working Memory', style='List Number')
    doc.add_paragraph('11. Guard → Evaluator: 收集 pre/runtime/post 各阶段指标', style='List Number')
    doc.add_paragraph('12. SDK → Agent: 返回 ToolCall 结果（含 guard 元数据）', style='List Number')

    doc.add_heading('3.2 Session 启动数据流', level=2)
    doc.add_paragraph('1. Agent 框架创建新 Session', style='List Number')
    doc.add_paragraph('2. AgentOS SDK 初始化: new AgentOS({...})', style='List Number')
    doc.add_paragraph('3. Memory Manager: 从 Semantic Memory 加载用户偏好 + 项目上下文', style='List Number')
    doc.add_paragraph('4. Memory Manager: 从 Episodic Memory 加载最近 7 天高重要性事件', style='List Number')
    doc.add_paragraph('5. Memory Manager: 生成上下文摘要（< 2KB）', style='List Number')
    doc.add_paragraph('6. SDK: 将上下文摘要注入 Agent 的 system prompt', style='List Number')
    doc.add_paragraph('7. Agent 开始正常工作，带完整上下文', style='List Number')

    # ====== 4. 接口规范 ======
    doc.add_heading('4. 模块间接口规范', level=1)

    doc.add_heading('4.1 Guard → Memory 接口', level=2)
    add_table(doc, ['接口', '方向', '数据', '频率'], [
        ['AuditLog → WM', 'Guard → WM', 'AuditEntry（完整）', '每次 tool call'],
        ['WM → EM', 'WM → EM（Session 结束时）', '标记为重要的事件', '每 Session'],
    ])

    doc.add_heading('4.2 Guard → Evaluator 接口', level=2)
    add_table(doc, ['接口', '方向', '数据', '频率'], [
        ['Schema Gate 结果', 'Guard → Eval', 'schema_pass, errors[]', '每次 tool call'],
        ['Risk Gate 结果', 'Guard → Eval', 'risk_score, action', '每次 tool call'],
        ['Verify Gate 结果', 'Guard → Eval', 'verify_status, checks[]', '每次 tool call'],
    ])

    # ====== 5. 安全设计 ======
    doc.add_heading('5. 安全设计', level=1)

    doc.add_heading('5.1 敏感数据保护', level=2)
    add_table(doc, ['机制', '说明'], [
        ['参数脱敏', '标记 x-secret=true 的参数在日志中替换为 [REDACTED]'],
        ['审计日志只读', 'JSONL 文件 IsReadOnly + 追加写入，禁止修改已写条目'],
        ['信任等级继承', '子 Agent 的信任等级 ≤ 父 Agent，防止提权'],
        ['操作溯源', '每条 AuditEntry 记录 sessionId + agentId，出事可定位'],
    ])

    doc.add_heading('5.2 防护层次', level=2)
    add_table(doc, ['层次', '防护内容', '实现'], [
        ['L1: 格式层', '参数类型/范围/格式', 'JSON Schema 校验（确定性）'],
        ['L2: 约束层', '路径范围/依赖关系/互斥', 'x- 扩展自定义校验（确定性）'],
        ['L3: 风险层', '影响×不可逆×敏感×错误率', '风险公式评分（确定性）'],
        ['L4: 验证层', '执行后结果校验', 'Verify Gate（确定性）'],
        ['L5: 审计层', '不可篡改操作记录', 'Append-only log（工程层）'],
    ])

    # ====== 6. 部署方案 ======
    doc.add_heading('6. 部署方案', level=1)

    doc.add_heading('6.1 MVP 单机部署', level=2)
    doc.add_paragraph('npm install @agentos/core')
    doc.add_paragraph('所有数据存储在本地：JSONL 日志 + JSON 记忆文件 + YAML 配置')
    doc.add_paragraph('Guard 规则通过 agentos.yaml 配置')
    doc.add_paragraph('适用：单个开发者的 Agent')

    doc.add_heading('6.2 v1.0 团队部署', level=2)
    doc.add_paragraph('npm install @agentos/core @agentos/server')
    doc.add_paragraph('SQLite 存储：共享 Audit Log + Memory')
    doc.add_paragraph('agentos-server 提供 Dashboard API')
    doc.add_paragraph('团队共享 Guard Rules（通过 Git 同步 agentos.yaml）')
    doc.add_paragraph('适用：3-10 人团队')

    doc.add_heading('6.3 v2.0 企业部署', level=2)
    doc.add_paragraph('npm install @agentos/cloud')
    doc.add_paragraph('远端 append-only Log Service（合规审计）')
    doc.add_paragraph('企业 Dashboard + RBAC 权限')
    doc.add_paragraph('SaaS 规则市场：社区共享 Guard Rules')
    doc.add_paragraph('适用：企业级、需要审计合规')

    doc.save(os.path.join(OUT_DIR, 'AgentOS_架构文档.docx'))
    print('[OK] AgentOS_架构文档.docx')


# ============================================================
# 文档 3: 任务清单 (TASKS.docx)
# ============================================================
def generate_tasks_doc():
    doc = Document()
    style_doc(doc)

    # 封面
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('AgentOS 任务清单')
    run.font.size = Pt(32)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('从 MVP 到 v2.0 的完整开发任务分解')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x88)

    doc.add_page_break()

    # ====== Phase 0: 项目初始化 ======
    doc.add_heading('Phase 0: 项目初始化', level=1)
    doc.add_paragraph('目标：搭建项目骨架，确保 CI/CD 和开发环境可用')
    doc.add_paragraph('预估时间：0.5 天')

    add_table(doc, ['#', '任务', '描述', '验收标准', '预估'], [
        ['0.1', '创建项目结构', '初始化 npm 包 @agentos/core，设置 TypeScript 严格模式', 'tsc --noEmit 通过', '1h'],
        ['0.2', '配置 ESLint + Prettier', '代码规范统一', 'eslint --quiet 通过', '0.5h'],
        ['0.3', '配置 Jest 测试框架', '单元测试 + 覆盖率报告', 'npx jest 运行通过 1 个示例测试', '0.5h'],
        ['0.4', '创建 GitHub 仓库', '创建 repo、配置 CI (GitHub Actions)', 'push → CI 绿灯', '1h'],
        ['0.5', '写 README.md', '项目说明、快速开始、架构概览', '新开发者照着 README 能跑起来', '1h'],
    ])

    # ====== Phase 1: Schema Gate ======
    doc.add_heading('Phase 1: Schema Gate（Guard 层第一关）', level=1)
    doc.add_paragraph('目标：实现对 tool call 参数的确定性校验，拦截格式错误')
    doc.add_paragraph('预估时间：3 天')

    add_table(doc, ['#', '任务', '描述', '验收标准', '预估'], [
        ['1.1', 'Tool Schema 类型定义', '定义 ToolSchema TypeScript 接口，含标准 JSON Schema + x- 扩展字段', '类型定义完整，x- 字段有 JSDoc', '2h'],
        ['1.2', 'Schema Registry', '工具注册表：add / get / list / remove', '注册 5 个测试工具，get 返回正确 schema', '1.5h'],
        ['1.3', 'JSON Schema 校验器', '集成 ajv，对 tool call 参数做标准 JSON Schema 校验', '无效参数→拦截；有效参数→放行。20 个测试用例', '3h'],
        ['1.4', 'Path Scope 校验器', '实现 x-path-scope 约束：workspace/temp/global 范围检查', '路径穿越攻击被拦截。10 个测试用例', '2h'],
        ['1.5', 'Path Allow/Deny 校验器', '实现 x-path-allow / x-path-deny Glob 模式匹配', '.git/ 路径被拦截。10 个测试用例', '2h'],
        ['1.6', 'Dependency 校验器', '实现 x-depends-on 参数依赖关系校验', 'auto_merge 无 base_branch → FAIL。5 个测试', '1.5h'],
        ['1.7', 'Mutual Exclusive 校验器', '实现 x-mutually-exclusive 互斥约束', 'content + file_path 同时存在 → FAIL。5 个测试', '1h'],
        ['1.8', 'Schema Gate 主入口', '整合所有校验器为单一入口：validate(toolCall) → GateResult', '30+ 测试用例全部通过', '2h'],
        ['1.9', 'Secret 脱敏', 'x-secret=true 的参数值替换为 [REDACTED]', '脱敏参数在 GateResult 中显示为 [REDACTED]', '1h'],
    ])

    # ====== Phase 2: Risk Gate ======
    doc.add_heading('Phase 2: Risk Gate（Guard 层第二关）', level=1)
    doc.add_paragraph('目标：实现四维风险评分，自动分级处置')
    doc.add_paragraph('预估时间：2 天')

    add_table(doc, ['#', '任务', '描述', '验收标准', '预估'], [
        ['2.1', '风险维度定义', '定义 Impact / Reversibility / Sensitivity 枚举和计算函数', '所有工具都有默认风险参数', '1.5h'],
        ['2.2', 'ErrorRate 追踪器', '记录每个工具的历史调用成功/失败统计', '调用 10 次后 ErrorRate 正确反映', '2h'],
        ['2.3', '风险评分引擎', '实现 RiskScore = Impact × Irrev × Sens × (1+ErrorRate)', '预定义的 10 个测试用例分数正确', '2h'],
        ['2.4', 'Action 分级器', '根据 RiskScore 和阈值决定 auto/notify/confirm/deny', '各阈值边界值测试通过', '1.5h'],
        ['2.5', 'Confirm 流程', '实现等待用户确认的机制（回调 + 超时）', '超时→拒绝；确认→执行；拒绝→取消', '2h'],
        ['2.6', 'Risk Gate 主入口', '整合评分 + 分级 + Confirm 流程', '完整流程测试通过', '1.5h'],
    ])

    # ====== Phase 3: Snapshot + Audit ======
    doc.add_heading('Phase 3: Snapshot + Audit Log', level=1)
    doc.add_paragraph('目标：实现执行前快照和不可篡改审计日志')
    doc.add_paragraph('预估时间：2 天')

    add_table(doc, ['#', '任务', '描述', '验收标准', '预估'], [
        ['3.1', 'File Hasher', '对指定文件/目录生成 sha256 hash', '同一文件 hash 一致，修改后 hash 不同', '1.5h'],
        ['3.2', 'Snapshot 创建器', '根据操作范围创建 Snapshot（fileHashes + gitHead + env）', 'Snapshot 包含所有被影响文件的 hash', '2h'],
        ['3.3', 'Snapshot 策略', '实现不同风险等级的快照策略（单文件/批量/全 workspace）', '高风险操作触发全量快照', '1h'],
        ['3.4', 'Audit Entry 构建器', '从 Guard 各阶段结果构建 AuditEntry', 'AuditEntry 包含所有阶段结果', '1.5h'],
        ['3.5', 'JSONL Logger', '追加写入 JSONL 文件，文件设为只读', '追加不覆盖；手动修改后校验失败', '2h'],
        ['3.6', 'Audit 查询器', '按时间范围/session/agent 查询 Audit Log', '查询结果正确，支持分页', '1.5h'],
    ])

    # ====== Phase 4: Verify Gate + Rollback ======
    doc.add_heading('Phase 4: Verify Gate + Rollback', level=1)
    doc.add_paragraph('目标：执行后校验结果，检测幻觉，支持一键回滚')
    doc.add_paragraph('预估时间：2 天')

    add_table(doc, ['#', '任务', '描述', '验收标准', '预估'], [
        ['4.1', 'File Exists 验证器', '工具声称创建/修改文件 → 检查文件是否存在', '不存在的文件→标记 FAIL', '1h'],
        ['4.2', 'File Changed 验证器', '工具声称修改文件 → 对比 Snapshot hash', 'hash 未变→标记 FAIL（幻觉）', '1h'],
        ['4.3', 'Lint 验证器', '工具修改代码文件 → 运行 eslint', 'lint 失败→标记 WARN', '1.5h'],
        ['4.4', 'TypeCheck 验证器', '工具修改 .ts 文件 → 运行 tsc --noEmit', 'type error→标记 WARN', '1h'],
        ['4.5', 'npm Verify 验证器', '工具声称发布 npm → 查询 registry 确认', 'registry 无新版本→标记 FAIL（幻觉）', '1h'],
        ['4.6', 'Git Push 验证器', '工具声称推送 → git ls-remote 确认', '远程无新 commit→标记 FAIL', '1h'],
        ['4.7', 'Verify Gate 主入口', '整合所有验证器，返回 PASS/WARN/FAIL', 'PASS=所有通过；WARN=非关键警告；FAIL=关键失败', '1.5h'],
        ['4.8', 'Rollback 引擎', '从 Snapshot 恢复文件状态', '回滚后文件 hash == Snapshot hash', '2h'],
        ['4.9', 'Verify + Rollback 联动', 'FAIL 操作自动触发回滚（可配置）', 'FAIL → 自动回滚 → 日志记录', '1h'],
    ])

    # ====== Phase 5: Guard Pipeline ======
    doc.add_heading('Phase 5: Guard 流水线整合', level=1)
    doc.add_paragraph('目标：将 Phase 1-4 整合为完整的 6 阶段 Guard Pipeline')
    doc.add_paragraph('预估时间：1 天')

    add_table(doc, ['#', '任务', '描述', '验收标准', '预估'], [
        ['5.1', 'Pipeline 编排器', '串联 6 阶段：Schema→Risk→Snapshot→Execute→Verify→Audit', '各阶段按序执行，Fail 正确中断', '3h'],
        ['5.2', 'Pipeline 配置化', '允许跳过/禁用特定阶段', '开发环境可跳过 Verify', '3h'],
        ['5.3', 'Guard.wrap() SDK', '一行代码包裹原始 executor', 'agentExecutor = aos.guard.wrap(orig)', '2h'],
        ['5.4', '端到端测试', '完整流程：恶意参数→拦截，正常参数→通过', '10个 E2E 场景通过', '2h'],
    ])

    # ====== Phase 6: Memory 层 ======
    doc.add_heading('Phase 6: Memory 层', level=1)
    doc.add_paragraph('目标：实现三层记忆架构（Working / Episodic / Semantic）')
    doc.add_paragraph('预估时间：4 天')

    add_table(doc, ['#', '任务', '描述', '验收标准', '预估'], [
        ['6.1', 'Working Memory 实现', '当前 session 消息+操作缓存，上下文预算管理', '超预算自动压缩旧消息', '3h'],
        ['6.2', 'EpisodicEvent 类型定义', '事件类型、重要性评分、压缩状态', '类型完整，含 JSDoc', '1h'],
        ['6.3', '重要性自动评分', '实现 Importance = Base × Recency × Freq × Feedback', '不同类型事件得分合理', '2h'],
        ['6.4', 'Episodic Store (SQLite)', '事件 CRUD + 时间范围查询 + 标签过滤', '查询 10万条事件 < 100ms', '3h'],
        ['6.5', '记忆压缩引擎', 'full→summary→one-liner→forgotten 自动流转', '旧事件自动压缩，重要事件永不压缩', '3h'],
        ['6.6', 'Semantic Store', '用户偏好、项目上下文、学到的规则', 'key-value + 版本追踪', '2h'],
        ['6.7', '规则提炼引擎', '同类 correction ≥3 次 → 自动提炼为 learnedRule', '规则自动生成并标记置信度', '3h'],
        ['6.8', 'Session 上下文注入', '启动时自动从 SM+EM 加载摘要注入', '注入内容 <2KB，含关键上下文', '3h'],
        ['6.9', 'Memory.recall() API', '语义+时间双维度检索', '搜"上周三改了什么"返回正确事件', '2h'],
        ['6.10', 'Memory 层测试', '各层单元测试 + 集成测试', '50+ 测试通过', '2h'],
    ])

    # ====== Phase 7: Evaluator ======
    doc.add_heading('Phase 7: Evaluator 层', level=1)
    doc.add_paragraph('目标：三阶段评估 + 隐性反馈捕获')
    doc.add_paragraph('预估时间：3 天')

    add_table(doc, ['#', '任务', '描述', '验收标准', '预估'], [
        ['7.1', 'Pre-exec Collector', '从 Schema Gate 收集指标', 'schema_pass_rate 正确计算', '2h'],
        ['7.2', 'Runtime Collector', '从 Execute 过程收集指标', 'retry_rate / timeout_rate', '2h'],
        ['7.3', 'Post-exec Collector', '从 Verify Gate + 用户行为收集', 'verify_pass_rate 正确', '2h'],
        ['7.4', '隐性反馈检测器', '实现 6 种隐性信号检测', 'accept_as_is / edit / redo 等', '3h'],
        ['7.5', 'Agent Score 计算', '综合指标 → 0-100 质量分', '分数随表现变化合理', '2h'],
        ['7.6', 'Report Generator', '生成评估报告（JSON + 趋势数据）', '24h 报告包含所有指标', '2h'],
        ['7.7', 'Eval 层测试', '各 Collector 单元测试', '30+ 测试通过', '2h'],
    ])

    # ====== Phase 8: SDK ======
    doc.add_heading('Phase 8: SDK 封装', level=1)
    doc.add_paragraph('目标：5 行代码接入，框架无关')
    doc.add_paragraph('预估时间：2 天')

    add_table(doc, ['#', '任务', '描述', '验收标准', '预估'], [
        ['8.1', 'AgentOS 主类', 'new AgentOS({guard, memory, eval}) 统一入口', '初始化所有模块', '2h'],
        ['8.2', 'Guard.wrap() API', '包裹任意 tool executor，一行接入', '拦截+校验+审计 全自动', '3h'],
        ['8.3', '配置加载器', '从 agentos.yaml 加载配置', '支持默认值+用户覆盖', '2h'],
        ['8.4', 'MCP Server 适配', '以 MCP Server 形式暴露能力', 'Claude/Copilot 可直接调用', '3h'],
        ['8.5', 'OpenClaw 适配', '原生支持 OpenClaw agent 接入', 'OpenClaw session 自动接入', '3h'],
        ['8.6', 'SDK 文档', 'API Reference + 快速开始 + 示例', '照着文档能接入', '3h'],
    ])

    # ====== Phase 9: 发布 ======
    doc.add_heading('Phase 9: 发布与运营', level=1)
    doc.add_paragraph('目标：发布 npm、GitHub、Product Hunt')
    doc.add_paragraph('预估时间：2 天')

    add_table(doc, ['#', '任务', '描述', '验收标准', '预估'], [
        ['9.1', 'npm 发布', '发布 @agentos/core v0.1.0', 'npm install 成功', '1h'],
        ['9.2', 'GitHub Release', '创建 Release + Release Notes', 'Release 页面完整', '1h'],
        ['9.3', 'Demo 项目', '用 AgentOS 包裹 OpenClaw agent 的完整 Demo', 'demo 可跑通', '3h'],
        ['9.4', 'Product Hunt 准备', 'Tagline、描述、截图、Logo', 'PH 提交就绪', '3h'],
        ['9.5', '社区文档', 'GitHub Discussions + Issue Templates', '社区可参与', '2h'],
    ])

    # ====== 路线图总结 ======
    doc.add_heading('总路线图', level=1)
    add_table(doc, ['阶段', '内容', '预估', '依赖'], [
        ['Phase 0', '项目初始化', '0.5 天', '无'],
        ['Phase 1', 'Schema Gate', '3 天', 'Phase 0'],
        ['Phase 2', 'Risk Gate', '2 天', 'Phase 1'],
        ['Phase 3', 'Snapshot + Audit', '2 天', 'Phase 1'],
        ['Phase 4', 'Verify + Rollback', '2 天', 'Phase 3'],
        ['Phase 5', 'Guard Pipeline', '1 天', 'Phase 1-4'],
        ['Phase 6', 'Memory 层', '4 天', 'Phase 5 (依赖 Audit Log)'],
        ['Phase 7', 'Evaluator 层', '3 天', 'Phase 5'],
        ['Phase 8', 'SDK 封装', '2 天', 'Phase 5-7'],
        ['Phase 9', '发布运营', '2 天', 'Phase 8'],
        ['', 'MVP 合计', '21.5 天', ''],
    ])

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('MVP 总工时约 21.5 个工作日（约 1 个月），单人全职。')
    run.bold = True

    doc.save(os.path.join(OUT_DIR, 'AgentOS_任务清单.docx'))
    print('[OK] AgentOS_任务清单.docx')


# ============================================================
if __name__ == '__main__':
    generate_design_doc()
    generate_architecture_doc()
    generate_tasks_doc()
    print('\n=== All 3 documents generated! ===')